import streamlit as st
import os
import sys
import cv2
import numpy as np
from PIL import Image
from ultralytics import YOLO
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import altair as alt
import time  # 用于控制帧率
from docx import Document
from docx.shared import Inches

# 禁止 ultralytics 库检查或下载模型
os.environ['ULTRALYTICS_NO_CHECKS'] = '1'

# 禁止 ultralytics 库检查或下载模型
os.environ['ULTRALYTICS_NO_CHECKS'] = '1'

# 添加指定文件夹到系统路径
ultralytics_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../.."))
sys.path.insert(0, ultralytics_dir)  # 插入到 sys.path 的最前面，确保优先加载

# 页面标题
st.title("🔬 血细胞检测")

# 侧边栏：模型选择和输入方式
with st.sidebar:
    st.markdown("### 🛠️ 检测配置")

    # 模型文件列表
    weight_dir = os.path.join(os.path.dirname(__file__), "../..", "weight")  # weight 文件夹路径
    model_files = {
        "yolov8n": os.path.join(weight_dir, "yolov8n.pt"),
        "YOLO-BC": os.path.join(weight_dir, "YOLO-BC.pt"),  # 暂时使用现有模型，后续可替换为血细胞专用模型
    }

    # 选择模型文件
    selected_model = st.selectbox("选择模型文件", list(model_files.keys()))

    def load_model(model_path):
        try:
            if not os.path.exists(model_path):  # 检查文件是否存在
                # raise FileNotFoundError(f"模型文件 {model_path} 不存在，请检查路径！")
                raise FileNotFoundError(f"模型文件不存在，请检查路径！")
            with st.spinner("正在加载模型..."):
                model = YOLO(model_path)  # 加载选择的模型
                # st.success(f"模型 {model_path} 加载成功！")
                st.success(f"模型加载成功！")
                return model
        except Exception as e:
            st.error(f"模型加载失败: {e}")
            return None


    model = load_model(model_files[selected_model])
    model_path = os.path.abspath(model_files[selected_model])

    # 选择输入方式
    input_type = st.radio("选择输入方式", ["上传图片", "上传视频", "本地摄像头拍照检测", "本地摄像头实时检测"])

    # 添加置信度滑块
    confidence_threshold = st.slider(
        "置信度阈值",
        min_value=0.0,
        max_value=1.0,
        value=0.25,  # 默认值
        step=0.01,
        help="调整置信度阈值以过滤检测结果。"
    )

# 初始化进度条
progress_bar = st.progress(0)

# 根据输入方式显示不同的界面
if input_type == "上传图片":
    st.markdown("### 📷 上传图片")

    # 支持批量上传文件
    uploaded_files = st.file_uploader("选择图片文件", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

    if uploaded_files:

        # 创建保存上传文件的文件夹
        if not os.path.exists("uploaded_images"):
            os.makedirs("uploaded_images")

        # 设置开始检测和重新检测按钮
        start_detection_button = st.button("开始检测")
        restart_detection_button = st.button("重新检测")

        # 记录批量处理的统计信息
        all_class_counts = {}
        all_confidences = []
        all_frame_data = []
        all_heatmap_data = []
        progress = st.progress(0)   # 初始化进度条
        total_file = len(uploaded_files)

        # 创建word文档
        doc = Document()
        doc.add_heading('血细胞检测结果报告', 0)

        # 在 Word 文档中添加原始图像

        # 批量上传图片
        if start_detection_button or restart_detection_button:
            if uploaded_files:
                for idx, uploaded_file in enumerate(uploaded_files):
                    # 加载图片
                    pil_image = Image.open(uploaded_file)
                    st.image(pil_image, caption=f"上传的图像: {uploaded_file.name}", use_container_width=True)

                    image = np.array(pil_image)
                    image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)

                    # 将原始图片存入 Word 文档
                    doc.add_paragraph(f"原始图像：{uploaded_file.name}")
                    original_image_path = f"uploaded_images/original_{uploaded_file.name}"
                    pil_image.save(original_image_path)  # 保存为文件后加载到 Word
                    doc.add_picture(original_image_path, width=Inches(4.0))

                    # 开始检测
                    with st.spinner(f"正在检测 {uploaded_file.name}..."):
                        # 使用 YOLOv8 进行目标检测
                        results = model(image, conf=confidence_threshold, imgsz=640)

                        # 获取检测结果
                        annotated_image = image
                        for result in results:
                            annotated_image = result.plot()  # 获取带检测框的图片

                            # 显示检测结果
                            st.image(annotated_image, caption=f"检测结果: {uploaded_file.name}", use_container_width=True)

                            # 输出检测到的目标信息
                            st.markdown("### 📊 检测到的血细胞：")
                            
                            # 统计各类血细胞数量
                            class_counts = {}
                            for box in result.boxes:
                                class_id = int(box.cls)  # 类别 ID
                                class_name = model.names[class_id]  # 类别名称
                                confidence = float(box.conf)  # 置信度
                                
                                # 更新类别计数
                                if class_name in class_counts:
                                    class_counts[class_name] += 1
                                else:
                                    class_counts[class_name] = 1
                            
                            # 显示血细胞类型及数量
                            cell_type_mapping = {
                                "RBC": "红细胞 (RBC)",
                                "WBC": "白细胞 (WBC)",
                                "PC": "血小板 (PC)",
                                "abnormal": "异常细胞"
                            }
                            
                            # 显示检测到的每种血细胞类型及数量
                            for class_name, count in class_counts.items():
                                display_name = cell_type_mapping.get(class_name, class_name)
                                st.write(f"- **{display_name}**: {count} 个")
                            
                            # 显示每个检测到的血细胞及其置信度
                            st.markdown("### 📊 检测详情：")
                            for box in result.boxes:
                                class_id = int(box.cls)  # 类别 ID
                                class_name = model.names[class_id]  # 类别名称
                                display_name = cell_type_mapping.get(class_name, class_name)
                                confidence = float(box.conf)  # 置信度
                                st.write(f"- **{display_name}** (置信度: {confidence:.4f})")

                            # 保存检测后的图片
                            annotated_image_bgr = cv2.cvtColor(annotated_image, cv2.COLOR_RGB2BGR)
                            annotated_path = f"uploaded_images/annotated_{uploaded_file.name}"
                            cv2.imwrite(annotated_path, annotated_image_bgr)

                            # 将检测后的图片添加到 Word 文档中
                            doc.add_paragraph(f"检测结果： {uploaded_file.name}")
                            doc.add_picture(annotated_path, width=Inches(4.0))

                            # 提供下载按钮
                            with open(annotated_path, "rb") as file:
                                st.download_button(
                                    label=f"下载检测后的图片 - {uploaded_file.name}",
                                    data=file,
                                    file_name=f"annotated_{uploaded_file.name}",
                                    mime="image/jpeg"
                                )



                            # 统计信息
                            for box in result.boxes:
                                class_name = model.names[int(box.cls)]
                                if class_name in all_class_counts:
                                    all_class_counts[class_name] += 1
                                else:
                                    all_class_counts[class_name] = 1

                                all_confidences.append(float(box.conf))

                            all_frame_data.append(len(result.boxes))
                    # 更新进度条
                    progress = int((idx + 1) / len(uploaded_files) * 100)
                    progress_bar.progress(progress)

        # 批量统计图表
        st.markdown("### 📊 血细胞统计分析")

        doc.add_heading("血细胞统计结果", level=1)

        # 计算各类血细胞数量
        rbc_count = all_class_counts.get("RBC", 0)  # 红细胞
        wbc_count = all_class_counts.get("WBC", 0)  # 白细胞
        platelets_count = all_class_counts.get("PC", 0)  # 血小板
        abnormal_count = all_class_counts.get("abnormal", 0)  # 异常细胞
        total_count = rbc_count + wbc_count + platelets_count + abnormal_count  # 总数

        # 显示血细胞统计结果
        st.markdown("#### 血细胞计数")
        st.markdown(f"**红细胞 (RBC)**: {rbc_count} 个")
        st.markdown(f"**白细胞 (WBC)**: {wbc_count} 个")
        st.markdown(f"**血小板 (Platelets)**: {platelets_count} 个")
        st.markdown(f"**异常细胞**: {abnormal_count} 个")
        st.markdown(f"**总计**: {total_count} 个")

        # 添加到Word文档
        doc.add_heading("血细胞计数", level=2)
        doc.add_paragraph(f"红细胞 (RBC): {rbc_count} 个")
        doc.add_paragraph(f"白细胞 (WBC): {wbc_count} 个")
        doc.add_paragraph(f"血小板 (Platelets): {platelets_count} 个")
        doc.add_paragraph(f"异常细胞: {abnormal_count} 个")
        doc.add_paragraph(f"总计: {total_count} 个")

        # 血细胞类别分布图
        class_df = pd.DataFrame(list(all_class_counts.items()), columns=["类别", "数量"])
        class_chart = alt.Chart(class_df).mark_bar().encode(
            x="数量:Q", y=alt.Y("类别:N", sort="-x"), text="数量:Q"
        ).properties(width=600, height=400)
        st.altair_chart(class_chart)

        # 保存 Word 文档并提供下载
        word_file_path = "detection_results.docx"
        doc.save(word_file_path)

        # 添加下载按钮
        with open(word_file_path, "rb") as doc_file:
            st.download_button(
                label="下载检测结果 Word 文档",
                data=doc_file,
                file_name="detection_results.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    # 清除结果按钮
    if st.button("清除结果"):
        st.rerun()


# 根据输入方式显示不同的界面
if input_type == "上传视频":
    st.markdown("### 🎥 上传视频")
    uploaded_file = st.file_uploader("选择一个视频", type=["mp4", "avi", "mov"])
    if uploaded_file is not None:
        # 将上传的文件保存为临时文件
        with open("temp_video.mp4", "wb") as f:
            f.write(uploaded_file.read())

        # 显示原始视频
        st.video("temp_video.mp4")

        # 点击按钮开始检测
        if st.button("开始检测"):
            try:
                # 打开视频文件
                cap = cv2.VideoCapture("temp_video.mp4")
                if not cap.isOpened():
                    st.error("无法打开视频文件，请检查文件格式或路径。")
                    st.stop()

                # 获取视频的帧率和总帧数
                fps = int(cap.get(cv2.CAP_PROP_FPS))
                total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

                # 创建视频写入对象
                output_file = "output_video.mp4"
                fourcc = cv2.VideoWriter_fourcc(*"mp4v")
                out = cv2.VideoWriter(output_file, fourcc, fps, (int(cap.get(3)), int(cap.get(4))))

                # 创建占位符用于显示检测结果
                video_placeholder = st.empty()

                # 初始化统计数据
                class_counts = {}  # 目标类别分布
                confidences = []  # 置信度分布
                frame_data = []  # 每一帧的目标数量
                heatmap_data = []  # 目标位置热力图数据

                # 逐帧处理视频
                for frame_idx in range(total_frames):
                    ret, frame = cap.read()
                    if not ret:
                        break

                    # 使用 YOLOv8 进行目标检测
                    results = model(frame, conf=confidence_threshold)

                    # 获取检测结果
                    for result in results:
                        # 绘制检测框和标签
                        annotated_frame = result.plot()  # 获取带检测框的帧

                        # 更新占位符内容
                        video_placeholder.image(annotated_frame, caption="检测结果", use_container_width=True)

                        # 将检测后的帧写入输出视频
                        out.write(annotated_frame)

                        # 统计目标类别和置信度
                        for box in result.boxes:
                            class_id = int(box.cls)
                            class_name = model.names[class_id]
                            confidence = float(box.conf)

                            # 更新类别分布
                            if class_name in class_counts:
                                class_counts[class_name] += 1
                            else:
                                class_counts[class_name] = 1

                            # 更新置信度分布
                            confidences.append(confidence)

                            # 更新目标位置热力图数据
                            x_center = (box.xyxy[0][0] + box.xyxy[0][2]) / 2
                            y_center = (box.xyxy[0][1] + box.xyxy[0][3]) / 2
                            heatmap_data.append((x_center, y_center))
                            
                        # 显示检测到的血细胞信息
                        st.markdown("### 📊 检测到的血细胞：")
                        
                        # 统计各类血细胞数量
                        detected_class_counts = {}
                        for box in result.boxes:
                            class_id = int(box.cls)  # 类别 ID
                            class_name = model.names[class_id]  # 类别名称
                            
                            # 更新类别计数
                            if class_name in detected_class_counts:
                                detected_class_counts[class_name] += 1
                            else:
                                detected_class_counts[class_name] = 1
                        
                        # 显示血细胞类型及数量
                        cell_type_mapping = {
                            "RBC": "红细胞 (RBC)",
                            "WBC": "白细胞 (WBC)",
                            "PC": "血小板 (PC)",
                            "abnormal": "异常细胞"
                        }
                        
                        # 显示检测到的每种血细胞类型及数量
                        for class_name, count in detected_class_counts.items():
                            display_name = cell_type_mapping.get(class_name, class_name)
                            st.write(f"- **{display_name}**: {count} 个")
                        
                        # 显示每个检测到的血细胞及其置信度
                        st.markdown("### 📊 检测详情：")
                        for box in result.boxes:
                            class_id = int(box.cls)  # 类别 ID
                            class_name = model.names[class_id]  # 类别名称
                            display_name = cell_type_mapping.get(class_name, class_name)
                            confidence = float(box.conf)  # 置信度
                            st.write(f"- **{display_name}** (置信度: {confidence:.2f})")


                        # 更新每一帧的目标数量
                        frame_data.append(len(result.boxes))

                    # 控制帧率（例如每秒处理 10 帧）
                    time.sleep(0.00008)  # 调整 sleep 时间以控制帧率

                # 释放视频文件
                cap.release()
                out.release()



                # 提供下载按钮
                with open(output_file, "rb") as file:
                    btn = st.download_button(
                        label="下载检测后的视频",
                        data=file,
                        file_name="output_video.mp4",
                        mime="video/mp4"
                    )

                # 血细胞统计分析
                st.markdown("### 📊 血细胞统计分析")

                # 计算各类血细胞数量
                rbc_count = class_counts.get("RBC", 0)  # 红细胞
                wbc_count = class_counts.get("WBC", 0)  # 白细胞
                platelets_count = class_counts.get("PC", 0)  # 血小板
                abnormal_count = class_counts.get("abnormal", 0)  # 异常细胞
                total_count = rbc_count + wbc_count + platelets_count + abnormal_count  # 总数

                # 显示血细胞统计结果
                st.markdown("#### 血细胞计数")
                st.markdown(f"**红细胞 (RBC)**: {rbc_count} 个")
                st.markdown(f"**白细胞 (WBC)**: {wbc_count} 个")
                st.markdown(f"**血小板 (Platelets)**: {platelets_count} 个")
                st.markdown(f"**异常细胞**: {abnormal_count} 个")
                st.markdown(f"**总计**: {total_count} 个")

                # 血细胞类别分布图
                st.markdown("#### 血细胞类别分布")
                class_df = pd.DataFrame(list(class_counts.items()), columns=["类别", "数量"])

                # 使用 Altair 创建水平柱状图
                class_chart = alt.Chart(class_df).mark_bar().encode(
                    x="数量:Q",
                    y=alt.Y("类别:N", sort="-x"),  # 按数量降序排列
                    text="数量:Q"
                ).properties(
                    width=600,
                    height=400
                )

                # 在柱体上方显示数值
                text = class_chart.mark_text(
                    align="left",
                    baseline="middle",
                    dx=3  # 数值与柱体的偏移量
                ).encode(
                    text="数量:Q"
                )

                st.altair_chart(class_chart + text)


            except Exception as e:
                st.error(f"视频处理失败: {e}")

        # 停止检测按钮
        if st.button("停止检测"):
            st.warning("检测已停止")


# 根据输入方式显示不同的界面
if input_type == "本地摄像头拍照检测":
    st.markdown("### 📸 本地摄像头拍照检测")
    # 使用 Streamlit 的 st.camera_input 调用本地摄像头
    camera_image = st.camera_input("本地摄像头拍照检测")

    if camera_image is not None:
        # 将摄像头画面转换为 OpenCV 格式
        file_bytes = np.asarray(bytearray(camera_image.read()), dtype=np.uint8)
        frame = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
        frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)  # 转换为 RGB 格式

        # 显示原始摄像头画面
        st.image(frame, caption="摄像头画面", use_container_width=True)

        # 点击按钮开始检测
        if st.button("开始检测"):
            with st.spinner("正在检测..."):
                # 使用 YOLOv8 进行目标检测
                results = model(frame, conf=confidence_threshold)

                # 获取检测结果
                for result in results:
                    # 绘制检测框和标签
                    annotated_frame = result.plot()  # 获取带检测框的帧

                    # 显示检测结果
                    st.image(annotated_frame, caption="检测结果", use_container_width=True)

                    # 输出检测到的目标信息
                    st.markdown("### 📊 检测到的血细胞：")
                    
                    # 统计各类血细胞数量
                    class_counts = {}
                    for box in result.boxes:
                        class_id = int(box.cls)  # 类别 ID
                        class_name = model.names[class_id]  # 类别名称
                        confidence = float(box.conf)  # 置信度
                        
                        # 更新类别计数
                        if class_name in class_counts:
                            class_counts[class_name] += 1
                        else:
                            class_counts[class_name] = 1
                    
                    # 显示血细胞类型及数量
                    cell_type_mapping = {
                        "RBC": "红细胞 (RBC)",
                        "WBC": "白细胞 (WBC)",
                        "PC": "血小板 (PC)",
                        "abnormal": "异常细胞"
                    }
                    
                    # 显示检测到的每种血细胞类型及数量
                    for class_name, count in class_counts.items():
                        display_name = cell_type_mapping.get(class_name, class_name)
                        st.write(f"- **{display_name}**: {count} 个")
                    
                    # 显示每个检测到的血细胞及其置信度
                    st.markdown("### 📊 检测详情：")
                    for box in result.boxes:
                        class_id = int(box.cls)  # 类别 ID
                        class_name = model.names[class_id]  # 类别名称
                        display_name = cell_type_mapping.get(class_name, class_name)
                        confidence = float(box.conf)  # 置信度
                        st.write(f"- **{display_name}** (置信度: {confidence:.2f})")

                    # 将检测后的图片保存为临时文件
                    annotated_frame_bgr = cv2.cvtColor(annotated_frame, cv2.COLOR_RGB2BGR)  # 转换回 BGR 格式
                    cv2.imwrite("annotated_frame.jpg", annotated_frame_bgr)

                    # 提供下载按钮
                    with open("annotated_frame.jpg", "rb") as file:
                        btn = st.download_button(
                            label="下载检测后的图片",
                            data=file,
                            file_name="annotated_frame.jpg",
                            mime="image/jpeg"
                        )

                    # 血细胞统计分析
                    st.markdown("### 📊 血细胞统计分析")

                    # 统计各类血细胞数量
                    class_counts = {}
                    for box in result.boxes:
                        class_id = int(box.cls)
                        class_name = model.names[class_id]
                        if class_name in class_counts:
                            class_counts[class_name] += 1
                        else:
                            class_counts[class_name] = 1

                    # 计算各类血细胞数量
                    rbc_count = class_counts.get("RBC", 0)  # 红细胞
                    wbc_count = class_counts.get("WBC", 0)  # 白细胞
                    platelets_count = class_counts.get("PC", 0)  # 血小板
                    abnormal_count = class_counts.get("abnormal", 0)  # 异常细胞
                    total_count = rbc_count + wbc_count + platelets_count + abnormal_count  # 总数

                    # 显示血细胞统计结果
                    st.markdown("#### 血细胞计数")
                    st.markdown(f"**红细胞 (RBC)**: {rbc_count} 个")
                    st.markdown(f"**白细胞 (WBC)**: {wbc_count} 个")
                    st.markdown(f"**血小板 (Platelets)**: {platelets_count} 个")
                    st.markdown(f"**异常细胞**: {abnormal_count} 个")
                    st.markdown(f"**总计**: {total_count} 个")

                    # 将类别分布数据转换为 DataFrame
                    class_df = pd.DataFrame(list(class_counts.items()), columns=["类别", "数量"])

                    # 使用 Altair 创建水平柱状图
                    class_chart = alt.Chart(class_df).mark_bar().encode(
                        x="数量:Q",
                        y=alt.Y("类别:N", sort="-x"),  # 按数量降序排列
                        text="数量:Q"
                    ).properties(
                        width=600,
                        height=400
                    )

                    # 在柱体上方显示数值
                    text = class_chart.mark_text(
                        align="left",
                        baseline="middle",
                        dx=3  # 数值与柱体的偏移量
                    ).encode(
                        text="数量:Q"
                    )

                    st.altair_chart(class_chart + text)

        # 清除结果按钮
        if st.button("清除结果"):
            st.rerun()


# 根据输入方式显示不同的界面
if input_type == "本地摄像头实时检测":
    st.markdown("### 📹 本地摄像头实时检测")
    
    # 创建一个占位符用于显示实时检测结果
    video_placeholder = st.empty()
    
    # 添加开始检测和停止检测按钮
    col1, col2 = st.columns(2)
    with col1:
        start_button = st.button("开始实时检测")
    with col2:
        stop_button = st.button("停止实时检测")
    
    # 创建一个占位符用于显示检测信息
    info_placeholder = st.empty()
    
    # 初始化统计数据
    class_counts = {}  # 目标类别分布
    confidences = []  # 置信度分布
    frame_data = []  # 每一帧的目标数量
    heatmap_data = []  # 目标位置热力图数据
    
    # 开始实时检测
    if start_button:
        try:
            # 打开摄像头
            cap = cv2.VideoCapture(0)  # 0 表示默认摄像头
            if not cap.isOpened():
                st.error("无法打开摄像头，请检查设备连接。")
                st.stop()
            
            # 获取摄像头的帧率和分辨率
            fps = int(cap.get(cv2.CAP_PROP_FPS))
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            # 创建视频写入对象（可选，用于保存检测结果）
            # output_file = "realtime_detection.mp4"
            # fourcc = cv2.VideoWriter_fourcc(*"mp4v")
            # out = cv2.VideoWriter(output_file, fourcc, fps, (width, height))
            
            # 设置会话状态变量用于控制检测循环
            st.session_state.detection_running = True
            
            # 帧计数器
            frame_count = 0
            
            # 实时检测循环
            while st.session_state.detection_running:
                # 读取一帧
                ret, frame = cap.read()
                if not ret:
                    st.error("无法读取摄像头画面。")
                    break
                
                # 每隔几帧进行一次检测（可以调整以控制性能）
                if frame_count % 3 == 0:  # 每3帧检测一次
                    # 使用 YOLOv8 进行目标检测
                    results = model(frame, conf=confidence_threshold)
                    
                    # 获取检测结果
                    for result in results:
                        # 绘制检测框和标签
                        annotated_frame = result.plot()  # 获取带检测框的帧
                        
                        # 更新占位符内容
                        video_placeholder.image(annotated_frame, caption="实时检测结果", use_container_width=True)
                        
                        # 可选：将检测后的帧写入输出视频
                        # out.write(annotated_frame)
                        
                        # 统计目标类别和置信度
                        detected_classes = []
                        for box in result.boxes:
                            class_id = int(box.cls)
                            class_name = model.names[class_id]
                            confidence = float(box.conf)
                            
                            # 记录检测到的类别
                            detected_classes.append(class_name)
                            
                            # 更新类别分布
                            if class_name in class_counts:
                                class_counts[class_name] += 1
                            else:
                                class_counts[class_name] = 1
                            
                            # 更新置信度分布
                            confidences.append(confidence)
                            
                            # 更新目标位置热力图数据
                            x_center = (box.xyxy[0][0] + box.xyxy[0][2]) / 2
                            y_center = (box.xyxy[0][1] + box.xyxy[0][3]) / 2
                            heatmap_data.append((x_center, y_center))
                        
                        # 更新每一帧的目标数量
                        frame_data.append(len(result.boxes))
                        
                        # 显示实时检测信息
                        info_text = f"### 📊 检测到的血细胞：\n"
                        
                        # 统计各类血细胞数量
                        detected_class_counts = {}
                        for cls in detected_classes:
                            if cls in detected_class_counts:
                                detected_class_counts[cls] += 1
                            else:
                                detected_class_counts[cls] = 1
                        
                        # 血细胞类型映射
                        cell_type_mapping = {
                            "RBC": "红细胞 (RBC)",
                            "WBC": "白细胞 (WBC)",
                            "PC": "血小板 (PC)",
                            "abnormal": "异常细胞"
                        }
                        
                        # 显示检测到的每种血细胞类型及数量
                        if detected_class_counts:
                            for class_name, count in detected_class_counts.items():
                                display_name = cell_type_mapping.get(class_name, class_name)
                                info_text += f"- **{display_name}**: {count} 个\n"
                        else:
                            info_text += "未检测到血细胞\n"
                        
                        info_placeholder.markdown(info_text)
                
                # 增加帧计数器
                frame_count += 1
                
                # 控制帧率
                time.sleep(0.01)  # 调整 sleep 时间以控制帧率
                
                # 检查是否点击了停止按钮
                if stop_button or not st.session_state.detection_running:
                    break
            
            # 释放资源
            cap.release()
            # out.release()
            
            # 显示检测结束信息
            st.success("实时检测已停止")
            
            # 如果有足够的数据，显示统计图表
            if frame_data and len(frame_data) > 5:
                st.markdown("### 📊 血细胞统计分析")
                
                # 计算各类血细胞数量
                rbc_count = class_counts.get("RBC", 0)  # 红细胞
                wbc_count = class_counts.get("WBC", 0)  # 白细胞
                platelets_count = class_counts.get("PC", 0)  # 血小板
                abnormal_count = class_counts.get("abnormal", 0)  # 异常细胞
                total_count = rbc_count + wbc_count + platelets_count + abnormal_count  # 总数

                # 显示血细胞统计结果
                st.markdown("#### 血细胞计数")
                st.markdown(f"**红细胞 (RBC)**: {rbc_count} 个")
                st.markdown(f"**白细胞 (WBC)**: {wbc_count} 个")
                st.markdown(f"**血小板 (Platelets)**: {platelets_count} 个")
                st.markdown(f"**异常细胞**: {abnormal_count} 个")
                st.markdown(f"**总计**: {total_count} 个")
                
                # 血细胞类别分布图
                st.markdown("#### 血细胞类别分布")
                class_df = pd.DataFrame(list(class_counts.items()), columns=["类别", "数量"])
                
                # 使用 Altair 创建水平柱状图
                class_chart = alt.Chart(class_df).mark_bar().encode(
                    x="数量:Q",
                    y=alt.Y("类别:N", sort="-x"),  # 按数量降序排列
                    text="数量:Q"
                ).properties(
                    width=600,
                    height=400
                )
                
                # 在柱体上方显示数值
                text = class_chart.mark_text(
                    align="left",
                    baseline="middle",
                    dx=3  # 数值与柱体的偏移量
                ).encode(
                    text="数量:Q"
                )
                
                st.altair_chart(class_chart + text)
        
        except Exception as e:
            st.error(f"实时检测失败: {e}")
    
    # 停止检测
    if stop_button:
        st.session_state.detection_running = False
        st.warning("检测已停止")